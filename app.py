import streamlit as st
import docx
import google.generativeai as genai
import json
from io import BytesIO

st.title("ADGM-Compliant Corporate Agent")

# ADGM Company Incorporation Document Checklist
ADGM_INCORPORATION_CHECKLIST = [
    "Articles of Association (AoA)",
    "Resolution of Board of Directors or Shareholders",
    "Certificate of Incorporation/Registration",
    "Passport and Signature Page",
    "UAE Visa Page",
    "Emirates ID",
    "UBO Declaration Form",
    "Business Plan",
    "Ownership and Organization Structure",
    "Human Resources and Physical Presence",
    "Financial Projections",
    "Consent to Act",
    "Nominee Declarations",
    "Data Protection Contact Person",
    "Economic Substance Form",
    "Source of Wealth Declaration Form",
    "Registered Office Address"
]

def check_missing_documents(uploaded_documents):
    """Checks for missing documents from the ADGM checklist."""
    uploaded_document_names = [doc.name for doc in uploaded_documents]
    missing_docs = []
    for item in ADGM_INCORPORATION_CHECKLIST:
        found = False
        for uploaded_name in uploaded_document_names:
            if item.lower() in uploaded_name.lower():
                found = True
                break
        if not found:
            missing_docs.append(item)
    return missing_docs

def analyze_document(document_content):
    """Analyzes the document content using the Gemini model."""
    prompt = f"""Analyze the following document and identify its type, any red flags, and suggest improvements based on ADGM regulations.

    Document Content:
    {document_content}

    Please provide the output in a JSON format with the following keys: 'document_type', 'red_flags', 'suggestions'.
    'red_flags' and 'suggestions' should be a list of dictionaries, each with 'section' and 'issue' keys.
    """
    model = genai.GenerativeModel('gemini-pro')
    response = model.generate_content(prompt)
    return json.loads(response.text)

def add_comments_to_doc(document, analysis_result):
    """Adds comments to the docx document based on the analysis."""
    for flag in analysis_result.get('red_flags', []):
        section_text = flag.get('section')
        issue = flag.get('issue')
        if section_text and issue:
            for para in document.paragraphs:
                if section_text in para.text:
                    para.add_comment(issue)
    return document

api_key = st.text_input("Enter your Google Generative AI API Key:", type="password")

uploaded_files = st.file_uploader("Upload .docx files", type=["docx"], accept_multiple_files=True)

if uploaded_files and api_key:
    st.write(f"{len(uploaded_files)} file(s) uploaded successfully!")

    missing_documents = check_missing_documents(uploaded_files)

    if missing_documents:
        st.warning("Missing Documents:")
        for doc in missing_documents:
            st.write(f"- {doc}")

    report = {
        "process": "Company Incorporation",
        "documents_uploaded": len(uploaded_files),
        "required_documents": len(ADGM_INCORPORATION_CHECKLIST),
        "missing_document": missing_documents,
        "issues_found": []
    }

    for uploaded_file in uploaded_files:
        try:
            genai.configure(api_key=AIzaSyDyHLhyLSOZx2ZASGqDfkgc38_v0V9gzMc)

            document = docx.Document(uploaded_file)
            content = ""
            for para in document.paragraphs:
                content += para.text + "\n"

            analysis_result = analyze_document(content)

            st.subheader(f"Analysis for: {uploaded_file.name}")
            st.json(analysis_result)

            report["issues_found"].append({
                "document": uploaded_file.name,
                "analysis": analysis_result
            })

            # Add comments to the document
            commented_doc = add_comments_to_doc(docx.Document(uploaded_file), analysis_result)

            # Save the commented document to a byte stream
            bio = BytesIO()
            commented_doc.save(bio)
            bio.seek(0)

            st.download_button(
                label=f"Download Reviewed {uploaded_file.name}",
                data=bio,
                file_name=f"reviewed_{uploaded_file.name}",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        except Exception as e:
            st.error(f"An error occurred with file {uploaded_file.name}: {e}")

    st.subheader("Final Report")
    st.json(report)

